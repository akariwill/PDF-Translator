import os, re
from translator.progress import Progress
from translator.llm_translator import LlmTranslator
from model import Model
from utils import LOG
from pdf2docx import Converter
from docx.shared import Cm
from docx import Document
from docx.enum.table import WD_ROW_HEIGHT_RULE

class DocParser:
    def __init__(self, model: Model, progress: Progress, target_language: str):
        self.translator = LlmTranslator(model, target_language)
        self.progress = progress
        self.taskConut = 0
        self.processed_paragraphs = set()  # Menyimpan hash dari paragraf yang sudah diproses

    def contains_email(text):
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b'
        return bool(re.search(email_pattern, text))

    def is_float(self, text):
        try:
            float(text)
            return True
        except ValueError:
            return False

    def process_text(self, text, is_count_mode: bool):
        if is_count_mode:
            self.taskConut += 1
            return
        else:
            self.progress.addCur()

        if text in ['\n', '']:
            return text
        if text.isnumeric() or self.is_float(text):
            return text

        translated_text = self.translator.translate_content(text)
        return translated_text

    def process_paragraph(self, paragraph, is_count_mode: bool):
        paragraph_hash = hash(paragraph.text)
        if paragraph_hash in self.processed_paragraphs:
            return
        self.processed_paragraphs.add(paragraph_hash)

        for run in paragraph.runs:
            if 'graphicData' in run._r.xml or 'oMathPara' in run._r.xml:
                return

        processed_text = self.process_text(paragraph.text, is_count_mode)
        if is_count_mode:
            return

        print("Hasil terjemahan:", processed_text)
        paragraph.clear()
        paragraph.add_run(processed_text)

    def process_all_paragraph(self, doc, is_count_mode: bool):
        if doc.paragraphs:
            for paragraph in doc.paragraphs:
                self.process_paragraph(paragraph, is_count_mode)
        return doc

    def process_table(self, table, is_count_mode: bool):
        for row in table.rows:
            row.height_rule = WD_ROW_HEIGHT_RULE.AUTO
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    self.process_paragraph(paragraph, is_count_mode)
                if cell.tables:
                    for nested_table in cell.tables:
                        self.process_table(nested_table, is_count_mode)

    def process_tables(self, doc, is_count_mode: bool):
        for table in doc.tables:
            self.process_table(table, is_count_mode)

    def convertDoc(self, pdf_input_path, convert_docx_file_path, endPos):
        cv = Converter(pdf_input_path)
        cv.convert(convert_docx_file_path, start=0, end=endPos)
        cv.close()

    def doTrans(self, file_name: str, result_docx_file_path: str, pdf_input_path: str, temp_source_path: str, endPos):
        endPos = endPos if endPos != 0 else None
        convert_docx_file_path = os.path.join(temp_source_path, file_name + ".docx")
        self.convertDoc(pdf_input_path, convert_docx_file_path, endPos)

        self.taskConut = 0
        self.progress.resetCur()
        doc = Document(convert_docx_file_path)

        self.process_tables(doc, True)
        self.process_all_paragraph(doc, True)
        self.progress.setAll(self.taskConut)

        self.processed_paragraphs = set()
        self.process_tables(doc, False)
        self.process_all_paragraph(doc, False)

        doc.save(result_docx_file_path)
        print('Proses selesai!')
